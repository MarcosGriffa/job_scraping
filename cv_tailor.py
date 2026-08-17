"""
cv_tailor.py — Genera un CV adaptado (.docx) por cada oferta finalista.

Reemplaza al cv_adapter.py viejo (hardcodeado 100% al CV de Marcos, con un
template de Node.js armado a mano). Esta versión:
  - Sirve para CUALQUIER CV (detecta las secciones automáticamente, sean
    del formato que sean — headers en mayúsculas normales o "espaciadas
    letra por letra" como las que deja PyMuPDF al extraer ciertos PDFs).
  - NO inventa nada: solo reordena y reformula el contenido que ya está en
    el CV real, para resaltar lo más relevante de cada oferta puntual.
    Cada reescritura de la IA pasa por una validación en Python (ver
    _is_safe_rewrite) que la descarta y usa el texto original si detecta
    palabras que no estaban en el CV.
  - Usa el modelo CHICO de Groq (ver ADAPT_MODEL) — es reescritura
    simple, no razonamiento pesado, así no depende del cupo del modelo
    grande (el que usa semantic_match.py para explicar matches).

Uso:
    python cv_tailor.py CV_Marcos_Griffa_v3.pdf --top 10
    python cv_tailor.py test_cvs/cv_ventas.txt --top 5

Guarda los .docx en data/cvs_adaptados/<nombre_cv>/
"""

from __future__ import annotations

import argparse
import html
import json
import os
import random
import re
import time
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Mm, Pt, RGBColor
from dotenv import load_dotenv
from groq import Groq

from cv_profile import extract_cv_text, classify_cv
from sources.computrabajo import fetch_full_description

load_dotenv()

# Modelo chico — reescritura simple, no razonamiento pesado. Cambiado el
# 17/08/2026: Groq dio de baja la línea Llama 3.x (ver cv_profile.py).
ADAPT_MODEL = "openai/gpt-oss-20b"

# Palabras clave (sin acentos, en mayúscula) que identifican una línea como
# título de sección de CV, sea cual sea el rubro o el idioma del CV.
SECTION_KEYWORDS = [
    "PERFIL", "RESUMEN", "OBJETIVO", "PROFILE", "SUMMARY",
    "PROYECTO", "PROJECT",
    "EXPERIENCIA", "EXPERIENCE",
    "EDUCACION", "FORMACION", "EDUCATION",
    "HABILIDAD", "SKILL", "COMPETENCIA", "APTITUD",
    "IDIOMA", "LANGUAGE",
    "CERTIFICACION", "CURSO", "CERTIFICATION",
    "SOBRE MI", "ABOUT ME",
    "CONTACTO", "CONTACT",
    "REFERENCIA", "REFERENCE",
    "LOGRO", "RECONOCIMIENTO",
    "INTERES", "HOBBIE", "HOBBY",
    "VOLUNTARIADO", "PUBLICACION",
]

# De esas secciones, cuáles tiene sentido adaptar por oferta (las demás —
# Educación, Idiomas, Sobre mí, etc. — se copian tal cual, sin tocarlas).
ADAPTABLE_KEYWORDS = [
    "PERFIL", "RESUMEN", "OBJETIVO", "PROFILE", "SUMMARY",
    "PROYECTO", "PROJECT",
    "EXPERIENCIA", "EXPERIENCE",
    "HABILIDAD", "SKILL", "COMPETENCIA", "APTITUD",
]

BULLET_MARKERS = ("-", "•", "●", "▪", "*", "·")


# ── Helpers de texto ─────────────────────────────────────────────
def _strip_accents(s: str) -> str:
    return "".join(c for c in unicodedata.normalize("NFKD", s) if not unicodedata.combining(c))


def _normalize_heading_text(raw: str) -> str:
    """Colapsa headers 'espaciados letra por letra' (ej. 'P E R F I L' → 'PERFIL')
    sin romper headers normales de dos palabras (ej. 'EXPERIENCIA LABORAL')."""
    raw = raw.strip()
    groups = re.split(r" {2,}", raw)  # 2+ espacios = separador de palabra real
    words = []
    for g in groups:
        tokens = [t for t in g.split(" ") if t]
        if tokens and all(len(t) == 1 for t in tokens):
            words.append("".join(tokens))  # letras sueltas de UNA palabra espaciada
        else:
            words.append(" ".join(tokens))
    return " ".join(w for w in words if w)


def _is_heading_line(raw_line: str) -> bool:
    stripped = raw_line.strip()
    if not stripped or len(stripped) > 45:
        return False
    if any(c.islower() for c in stripped):
        return False
    if not any(c.isalpha() for c in stripped):
        return False
    clean = _strip_accents(_normalize_heading_text(stripped)).upper()
    return any(kw in clean for kw in SECTION_KEYWORDS)


def _section_kind(heading: str) -> str:
    h = _strip_accents(heading).upper()
    if any(k in h for k in ("PERFIL", "RESUMEN", "OBJETIVO", "PROFILE", "SUMMARY")):
        return "perfil"
    if any(k in h for k in ("PROYECTO", "PROJECT")):
        return "proyectos"
    if any(k in h for k in ("EXPERIENCIA", "EXPERIENCE")):
        return "experiencia"
    if any(k in h for k in ("HABILIDAD", "SKILL", "COMPETENCIA", "APTITUD")):
        return "habilidades"
    return "other"


def _is_company_line(line: str) -> bool:
    """Detecta líneas de 'nombre de empresa' dentro de Experiencia laboral,
    tolerando que le siga una fecha en texto (ej. 'TIENDA XXI (03/2022 - Actualidad)')."""
    s = line.strip()
    if not s or len(s) > 80 or s[0] in BULLET_MARKERS:
        return False
    core = re.sub(r"\s*\([^)]*\)\s*$", "", s).strip()  # saca un "(...)" final si hay
    if not core or not any(c.isalpha() for c in core):
        return False
    return not any(c.islower() for c in core)


_HARD_STOP_CHARS = (".", ":", ")", "?", "!", ";")
_DANGLING_CHARS = ("·", "-", ",", "/", "&")
_DANGLING_WORDS = {"y", "o", "e", "u", "de", "del", "la", "el", "en", "con"}


def _ends_with_dangling(line: str) -> bool:
    stripped = line.rstrip()
    if not stripped:
        return False
    if stripped[-1] in _DANGLING_CHARS:
        return True
    last_word = re.split(r"\s+", stripped)[-1].strip(".,;:").lower()
    return last_word in _DANGLING_WORDS


def _rejoin_wrapped_lines(lines: list[str]) -> list[str]:
    """Reune líneas que en el PDF original quedaron cortadas a mitad de
    frase por el ancho angosto de una columna/sidebar, y las vuelve a unir
    en un solo párrafo fluido.

    Sin esto, texto como "Juego al rugby desde los 10 años," seguido de
    "deporte que me formó en disciplina," (dos líneas físicas separadas en
    el PDF, pero UNA sola oración) se renderiza en el .docx como dos
    párrafos sueltos con espacio entre medio — se ve "cortado"/roto (bug
    real encontrado 11/08/2026 en las secciones que se copian tal cual,
    como Educación y Sobre mí).

    Heurística: una línea es continuación de la anterior si la anterior
    NO termina en puntuación de cierre (. : ) ? ! ;) Y además la línea
    actual arranca en minúscula, o la anterior termina "colgando" (ej.
    "Estadística · Bases de Datos ·"), o la línea actual arranca con "(" .
    """
    paragraphs: list[str] = []
    for raw in lines:
        line = re.sub(r"\s+", " ", raw.strip())
        if not line:
            continue
        is_continuation = False
        if paragraphs:
            prev = paragraphs[-1]
            ends_hard_stop = bool(prev) and prev[-1] in _HARD_STOP_CHARS
            if not ends_hard_stop:
                if line[0].islower():
                    is_continuation = True
                elif _ends_with_dangling(prev):
                    is_continuation = True
                elif line.startswith("("):
                    is_continuation = True
        if is_continuation:
            paragraphs[-1] = f"{paragraphs[-1]} {line}"
        else:
            paragraphs.append(line)
    return paragraphs


def _content_words(text: str) -> set[str]:
    stopwords = {
        "para", "como", "desde", "hasta", "donde", "cuando", "porque", "tambien",
        "entre", "sobre", "este", "esta", "estos", "estas", "muy", "mas", "pero",
        "segun", "dentro", "fuera", "cada", "todo", "toda", "todos", "todas",
        "tiene", "tienen", "siendo", "siempre", "ademas", "otros", "otras", "with",
        "that", "this", "from", "your", "have", "will",
    }
    words = re.findall(r"[A-Za-zÁÉÍÓÚÑáéíóúñ]{4,}", text or "")
    return {w2 for w in words if (w2 := _strip_accents(w).lower()) not in stopwords}


def _is_safe_rewrite(cv_text: str, candidate: str, max_new_words: int = 0) -> bool:
    """Guarda anti-invención: si la reescritura de la IA trae palabras de
    contenido que NO estaban en ningún lado del CV original, la rechazamos
    y en su lugar se usa el texto original sin tocar.

    Tolerancia CERO por defecto. Probado con margen de 1 palabra y no
    alcanzó: en una corrida real cambió "Estudiante de Ciencia de Datos"
    por "Analista de datos" en el perfil de Marcos — un solo cambio de
    palabra, pero que le inventa un estatus laboral que no tiene. Con "no
    inventar nada" de por medio, un CV no es el lugar para admitir ningún
    margen de duda: se prefiere perder algo de reformulación creativa
    (queda el texto original, que reordenar/acortar igual sigue permitido)
    antes que arriesgarse a una sola palabra que no sea 100% cierta."""
    if not candidate or not candidate.strip():
        return False
    cand_words = _content_words(candidate)
    if not cand_words:
        return True
    original_pool = _content_words(cv_text)
    new_words = cand_words - original_pool
    return len(new_words) <= max_new_words


def _safe_filename(s: str, max_len: int = 40) -> str:
    s = _strip_accents(s or "")
    s = re.sub(r"[^A-Za-z0-9 _-]", "", s)
    s = re.sub(r"\s+", "_", s.strip())
    return s[:max_len]


# ── Parseo del CV en secciones ───────────────────────────────────
def parse_cv(cv_text: str) -> tuple[list[str], list[dict]]:
    """Separa el CV en el bloque de encabezado (nombre + contacto) y una
    lista de secciones {heading, adaptable, body_lines}, en el mismo orden
    en que aparecen en el CV original."""
    header_lines: list[str] = []
    sections: list[dict] = []
    current = None

    for raw_line in cv_text.split("\n"):
        line = raw_line.rstrip()
        if _is_heading_line(line):
            heading = _normalize_heading_text(line)
            adaptable = any(kw in _strip_accents(heading).upper() for kw in ADAPTABLE_KEYWORDS)
            current = {"heading": heading, "adaptable": adaptable, "body_lines": []}
            sections.append(current)
            continue
        if current is None:
            if line.strip():
                header_lines.append(line.strip())
        else:
            current["body_lines"].append(line)

    return header_lines, sections


def parse_experience_blocks(body_lines: list[str]) -> list[dict]:
    """Dentro de la sección de Experiencia laboral, separa por empresa:
    {company, meta_lines (rol/fechas, no se tocan), bullets (sí se adaptan)}.

    Ojo con los PDF: a veces un bullet largo se corta en dos líneas físicas
    (la marca "-" queda sola, y encima el texto puede seguir en una TERCERA
    línea sin marca) — mientras estemos "dentro" de un bullet, cualquier
    línea sin marca nueva ni empresa nueva se pega como continuación del
    bullet anterior, no como un meta_line suelto."""
    blocks: list[dict] = []
    current = None
    in_bullet = False  # true mientras las líneas siguientes son continuación del último bullet

    for raw in body_lines:
        line = raw.strip()
        if not line:
            continue
        if _is_company_line(line):
            current = {"company": line, "meta_lines": [], "bullets": []}
            blocks.append(current)
            in_bullet = False
            continue
        if current is None:
            continue  # texto suelto antes de la primera empresa detectada
        if line in BULLET_MARKERS:
            current["bullets"].append("")
            in_bullet = True
            continue
        if line[0] in BULLET_MARKERS and len(line) > 1:
            current["bullets"].append(line[1:].strip(" -•●▪*·"))
            in_bullet = True
            continue
        if in_bullet and current["bullets"]:
            current["bullets"][-1] = (current["bullets"][-1] + " " + line).strip()
        else:
            current["meta_lines"].append(line)

    return blocks


# ── Adaptación con IA (modelo chico) ─────────────────────────────
ADAPT_SYSTEM_PROMPT = """\
Sos un asesor de carrera. Vas a recibir fragmentos REALES de un CV y una
oferta laboral. Tu trabajo es adaptar la REDACCIÓN Y EL ORDEN de esos
fragmentos para que resalten lo más relevante para ESA oferta puntual.

REGLAS ABSOLUTAS:
- NO inventes experiencia, proyectos, habilidades, empresas, fechas, cifras
  ni tecnologías que no estén ya en el material que te paso.
- NO agregues ningún tipo de venta, técnica, herramienta, adjetivo o
  cualidad que no aparezca YA escrita en el material (ej: si no dice
  "ventas consultivas" o "negociación", vos tampoco lo escribas, aunque te
  parezca implícito o típico del rubro).
- Podés reordenar (lo más relevante primero) o acortar, pero preferí
  reusar las mismas palabras y frases del original antes que reformular
  con sinónimos — cuanto más textual, mejor. Ante la duda, dejá la frase
  como está en vez de arriesgarte a agregar algo que no está.
- Si te paso una lista de habilidades ("cv_habilidades"), devolvé la MISMA
  lista reordenada — ni un ítem distinto, ni de más ni de menos.
- Si te paso bullets de experiencia de una empresa ("cv_experiencia"),
  devolvé la MISMA CANTIDAD de bullets para esa empresa, reordenados y
  opcionalmente reformulados, pero sin agregar tareas nuevas.
- Si un campo no te lo paso, no lo incluyas en la respuesta.

Devolvé ÚNICAMENTE un JSON con este esquema:
{
  "perfil_adaptado": "string",
  "proyectos_adaptado": "string",
  "habilidades_adaptado": ["lista", "de", "strings"],
  "experiencia_adaptado": {"NOMBRE EMPRESA": ["bullet 1", "bullet 2"]}
}
"""


def _get_client() -> Groq:
    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        raise RuntimeError("Falta GROQ_API_KEY en el .env.")
    return Groq(api_key=api_key)


def adapt_for_job(cv_text: str, sections: list[dict], job: dict) -> dict:
    """Arma el payload con lo adaptable del CV, se lo manda al modelo chico,
    y valida cada campo de la respuesta antes de usarlo (si no pasa la
    validación, se usa el contenido original tal cual — nunca se inventa)."""
    payload: dict = {}

    for sec in sections:
        if not sec["adaptable"]:
            continue
        kind = _section_kind(sec["heading"])

        if kind == "perfil" and "cv_perfil" not in payload:
            payload["cv_perfil"] = "\n".join(_rejoin_wrapped_lines(sec["body_lines"]))
        elif kind == "proyectos" and "cv_proyectos" not in payload:
            payload["cv_proyectos"] = "\n".join(_rejoin_wrapped_lines(sec["body_lines"]))
        elif kind == "habilidades" and "cv_habilidades" not in payload:
            payload["cv_habilidades"] = [l.strip() for l in sec["body_lines"] if l.strip()]
        elif kind == "experiencia" and "cv_experiencia" not in payload:
            blocks = parse_experience_blocks(sec["body_lines"])
            sec["_blocks"] = blocks  # lo reusa build_docx para renderizar
            payload["cv_experiencia"] = {b["company"]: b["bullets"] for b in blocks if b["bullets"]}

    if not any(k.startswith("cv_") for k in payload):
        return {}  # no se detectó ninguna sección adaptable en este CV

    payload["oferta_titulo"] = job.get("title", "")
    payload["oferta_empresa"] = job.get("company", "")
    payload["oferta_descripcion"] = (job.get("description") or "")[:2000]

    client = _get_client()
    raw_json = None
    for attempt in range(2):
        try:
            resp = client.chat.completions.create(
                model=ADAPT_MODEL,
                messages=[
                    {"role": "system", "content": ADAPT_SYSTEM_PROMPT},
                    {"role": "user", "content": json.dumps(payload, ensure_ascii=False)},
                ],
                temperature=0.3,
                response_format={"type": "json_object"},
            )
            raw_json = json.loads(resp.choices[0].message.content.strip())
            break
        except Exception as e:
            print(f"    [cv_tailor] intento {attempt + 1} de adaptación falló: {e}")
    if raw_json is None:
        return {}

    result: dict = {}

    if "cv_perfil" in payload:
        cand = raw_json.get("perfil_adaptado", "")
        result["perfil"] = cand if _is_safe_rewrite(cv_text, cand) else payload["cv_perfil"]

    if "cv_proyectos" in payload:
        cand = raw_json.get("proyectos_adaptado", "")
        result["proyectos"] = cand if _is_safe_rewrite(cv_text, cand) else payload["cv_proyectos"]

    if "cv_habilidades" in payload:
        cand = raw_json.get("habilidades_adaptado", [])
        original_norm = sorted(_strip_accents(s).lower().strip() for s in payload["cv_habilidades"])
        cand_norm = sorted(_strip_accents(s).lower().strip() for s in cand) if isinstance(cand, list) else []
        result["habilidades"] = cand if cand_norm == original_norm else payload["cv_habilidades"]

    if "cv_experiencia" in payload:
        exp_out = {}
        cand_map = raw_json.get("experiencia_adaptado", {})
        cand_map = cand_map if isinstance(cand_map, dict) else {}
        cand_map_norm = {_strip_accents(k).upper().strip(): v for k, v in cand_map.items()}
        for company, original_bullets in payload["cv_experiencia"].items():
            cand_bullets = cand_map_norm.get(_strip_accents(company).upper().strip())
            if (
                isinstance(cand_bullets, list)
                and len(cand_bullets) == len(original_bullets)
                and all(isinstance(b, str) and _is_safe_rewrite(cv_text, b) for b in cand_bullets)
            ):
                exp_out[company] = cand_bullets
            else:
                exp_out[company] = original_bullets  # fallback seguro para ESTA empresa
        result["experiencia"] = exp_out

    return result


# ── Render a .docx — template visual (referencia: Canva de Marcos) ──
#
# Header navy de ancho completo + barra dorada de acento + cuerpo a dos
# columnas (tabla 1x2 sin bordes). No tenemos el archivo fuente original
# de diseño (Canva), así que esto es una reconstrucción manual a partir de
# la descripción del usuario — queda como "template default" único por
# ahora. El contenido (texto) sigue pasando por las mismas reglas de
# "cero invención" de siempre; esto es 100% presentación visual.
NAVY = "1A1D2E"
GOLD = "F2C230"
GRAY_LIGHT = "E8E8E8"
GRAY_TEXT = "3A3A3A"
GRAY_ON_NAVY = "C9CBD6"
WHITE = "FFFFFF"

_SKILL_SUBHEADER_KEYWORDS = {"IDIOMAS", "LANGUAGES", "COMPETENCIAS", "APTITUDES"}

# ── Escala dinámica de tamaño/espaciado ──────────────────────────
#
# Con un CV cargado (4 proyectos, 2 experiencias, ~19 skills) el cuerpo ya
# ocupa toda la página con tamaños "normales". Pero con menos contenido
# (ej. la IA recortó a 1 proyecto y acortó bullets), esos mismos tamaños
# dejaban mucho aire abajo de la página — a pedido (11/08/2026), en vez de
# dejar el hueco, estiramos el contenido para que ocupe la hoja completa.
# El header y la barra dorada NUNCA escalan — solo el cuerpo.
#
# DOS escalas independientes, no una sola (lección real, 11/08/2026):
# agrandar la FUENTE tiene techo bajo — cada punto de más hace que las
# líneas envuelvan más seguido, y ese envolvido extra crece la altura de
# forma no lineal hasta desbordar a una 2da página (probado: escala 1.8
# combinada ya desborda un CV corto). El ESPACIADO (entre líneas,
# secciones, píldoras) en cambio es lineal — no genera líneas nuevas,
# así que ahí se puede empujar mucho más fuerte sin ese riesgo. Por eso
# _FONT_SCALE tiene un techo bajo y _SPACE_SCALE uno mucho más alto.
_FONT_SCALE = 1.0
_SPACE_SCALE = 1.0


def _set_scale(font_scale: float, space_scale: float):
    global _FONT_SCALE, _SPACE_SCALE
    _FONT_SCALE = font_scale
    _SPACE_SCALE = space_scale


def _pt(value: float) -> Pt:
    return Pt(value * _FONT_SCALE)


def _content_score(cv_data: dict) -> int:
    """Estimación gruesa de 'cuánto texto hay' para decidir la escala.
    No mide layout real (eso solo lo sabe Word) — son pesos calibrados a
    ojo comparando contra el caso más cargado de Marcos, que a escala 1.0
    llena la página justo. Los "+N" fijos son porque cada bullet/píldora
    también consume espacio por su propio salto de línea/padding, no solo
    por el largo del texto."""
    score = 0
    for sec in cv_data.get("left_sections", []) + cv_data.get("right_sections", []):
        kind, content = sec["kind"], sec["content"]
        if kind in ("perfil", "generic"):
            score += sum(len(t) for t in content)
        elif kind == "proyectos":
            score += sum(len(item["text"]) for item in content)
        elif kind == "experiencia":
            for block in content:
                score += len(block["company"])
                score += sum(len(m) for m in block["meta_lines"])
                score += sum(len(b) + 15 for b in block["bullets"])
        elif kind == "habilidades":
            score += sum(len(item["text"]) + 25 for item in content["core"])
            for group in content["groups"]:
                score += len(group["label"]) + len(group["text"])
    return score


def _pick_scales(cv_data: dict) -> tuple[float, float]:
    """Devuelve (font_scale, space_scale). Umbrales calibrados contra el
    CV real de Marcos sin recortar por la IA (score ~2900-3300, llena la
    página a escala 1.0/1.0) y probando a mano con una versión mucho más
    corta hasta encontrar el techo de espaciado que no desborda a una 2da
    página — probado a mano (11/08/2026): score 2376 aguanta space_scale
    1.8 pero NO 2.2; score 1407 aguanta 4.0 pero no 4.5; score 1152
    aguanta 4.5 pero no 5.5. Los valores de abajo quedan con margen por
    debajo de esos techos medidos (un CV real dentro del mismo escalón
    puede tener un poco más de texto que el caso que probé)."""
    score = _content_score(cv_data)
    if score >= 2400:
        return 1.0, 1.0
    if score >= 1800:
        return 1.04, 1.6
    if score >= 1300:
        return 1.12, 3.5
    if score >= 900:
        return 1.16, 4.0
    return 1.2, 4.2  # piso de fuente: pasado este punto ya se ve raro en un CV


def _tight(paragraph, space_before: float = 0, space_after: float = 4):
    """Espaciado compacto — sin esto, el interlineado por defecto de Word
    hace que un CV con harto contenido (como el de Marcos: 4 proyectos, 2
    experiencias, ~15 skills) se vaya a varias páginas sin necesidad.

    space_before/after escalan con _SPACE_SCALE (el que puede llegar a
    3x+ en un CV corto, sin riesgo de desbordar — ver nota arriba de por
    qué está separado de _FONT_SCALE). El interlineado DENTRO de un mismo
    párrafo (line_spacing) también ayuda a repartir el contenido, pero se
    lo deja con un techo bajo (1.3 máx) — a diferencia del espacio ENTRE
    párrafos, un interlineado grande dentro de una misma oración se ve
    raro (como doble espacio en un word processor de los 90)."""
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(space_before * _SPACE_SCALE)
    fmt.space_after = Pt(space_after * _SPACE_SCALE)
    fmt.line_spacing = min(1.0 + (_SPACE_SCALE - 1.0) * 0.12, 1.3)


def _hex_to_rgb(hex_color: str) -> RGBColor:
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def _shade_element(el, hex_color: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    el.append(shd)


def _shade_cell(cell, hex_color: str):
    _shade_element(cell._tc.get_or_add_tcPr(), hex_color)


def _shade_run(run, hex_color: str):
    _shade_element(run._element.get_or_add_rPr(), hex_color)


def _set_run_spacing(run, points: float):
    """Letter-spacing — python-docx no lo expone como propiedad, se inyecta
    a mano en el XML (w:spacing, en veinteavos de punto)."""
    rPr = run._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), str(int(points * 20)))
    rPr.append(spacing)


def _set_paragraph_bottom_border(paragraph, hex_color: str, size: int = 6, space: int = 2):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _remove_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


def _set_col_widths(table, widths):
    """widths: lista de Length (ej. Emu(...)/Cm(...)) — uno por columna."""
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width
    tblGrid = table._tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        for grid_col, width in zip(tblGrid.findall(qn("w:gridCol")), widths):
            grid_col.set(qn("w:w"), str(width.twips))


def _set_cell_margins(cell, top: int = 20, bottom: int = 20, left: int = 80, right: int = 80):
    """Márgenes internos de una celda, en veinteavos de punto (dxa).

    OJO con el orden: el schema de w:tcMar exige top, left, bottom, right
    en ESE orden exacto — con "bottom" antes que "left" (como estaba antes)
    Word ignoraba el elemento entero silenciosamente y usaba el margen por
    default de la tabla, así que todos los ajustes de padding de las
    píldoras no hacían nada (bug real encontrado 11/08/2026)."""
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tcPr.append(mar)


def _fill_pill_cell(outer_cell, text: str, hex_bg: str, hex_text: str, bold: bool):
    """Llena UNA celda de la grilla de píldoras (ver _add_skill_pill_grid)
    con una mini-tabla de 1x1 shrink-wrapped al texto — la celda de la
    grilla le da la posición ("fila, columna"), la mini-tabla de adentro
    le da el efecto píldora chica (fondo ajustado al texto, no una barra
    que ocupa toda la celda)."""
    _tight(outer_cell.paragraphs[0], space_before=0, space_after=0)
    pill_table = outer_cell.add_table(rows=1, cols=1)
    pill_table.autofit = True
    _remove_table_borders(pill_table)
    cell = pill_table.rows[0].cells[0]
    _shade_cell(cell, hex_bg)
    _set_cell_margins(cell, top=8, bottom=8, left=90, right=90)
    p = cell.paragraphs[0]
    _tight(p, space_before=0, space_after=0)
    run = p.add_run(text)
    run.font.size = _pt(9)
    run.font.color.rgb = _hex_to_rgb(hex_text)
    run.bold = bold
    # add_table() ya insertó solo un <w:p> DENTRO de esta celda de grilla
    # (mismo patrón que dentro de la mini-tabla) — lo achicamos también.
    _tight(outer_cell.paragraphs[-1], space_before=0, space_after=0)


def _add_skill_pill_grid(parent_cell, items: list[tuple[str, str, str, bool]], cols: int = 2):
    """Píldoras reales, VARIAS POR FILA — no una por línea ocupando todo
    el ancho de la columna (así se veía antes: una barra apilada tras
    otra, reportado como bug visual el 11/08/2026, no lo que pide el
    diseño de Canva). Se logra con una grilla de `cols` columnas sin
    bordes ni sombreado propio (solo da la estructura fila/columna); cada
    celda de esa grilla tiene adentro su propia mini-tabla-píldora
    shrink-wrapped al texto (ver _fill_pill_cell).

    items: lista de (texto, color_fondo, color_texto, negrita)."""
    if not items:
        return
    rows = (len(items) + cols - 1) // cols
    grid = parent_cell.add_table(rows=rows, cols=cols)
    grid.autofit = True
    _remove_table_borders(grid)
    idx = 0
    for r in range(rows):
        for c in range(cols):
            cell = grid.rows[r].cells[c]
            if idx < len(items):
                text, hex_bg, hex_text, bold = items[idx]
                _fill_pill_cell(cell, text, hex_bg, hex_text, bold)
            else:
                _tight(cell.paragraphs[0], space_before=0, space_after=0)
            idx += 1
    # add_table() ya insertó solo un <w:p> DESPUÉS de la grilla entera —
    # lo reusamos como separador antes del siguiente contenido.
    _tight(parent_cell.paragraphs[-1], space_before=0, space_after=6)


def _next_para(cell):
    """Las celdas de una tabla arrancan con un párrafo vacío ya puesto —
    lo reusamos la primera vez en vez de dejarlo como una línea en blanco
    de más antes del contenido real."""
    p = cell.paragraphs[-1]
    if p.runs or p.text:
        p = cell.add_paragraph()
    return p


def _contact_icon(line: str) -> str:
    if "@" in line:
        return "✉  "
    low = line.lower()
    if "linkedin" in low or "github" in low:
        return "▪  "
    if re.search(r"\d{2,}", line):
        return "☎  "
    return "▪  "


def _looks_like_title(text: str) -> bool:
    """Heurística para distinguir el nombre de un proyecto (línea corta,
    sin punto final) de su descripción (oración completa)."""
    t = text.strip()
    return bool(t) and len(t) <= 90 and t[-1] not in ".!?"


def _build_skill_pills(skills: list[str], top_k: int = 4) -> list[dict]:
    """Arma el contenido de Habilidades: las skills técnicas "core" (antes
    del primer subgrupo) se muestran como píldoras individuales — las
    primeras top_k, después del reordenamiento por IA si lo hubo, en
    dorado (son las que la IA priorizó como más relevantes para ESTA
    oferta puntual), el resto en gris.

    Los subgrupos tipo "Idiomas"/"Competencias" (e Idiomas/Competencias en
    sí mismos NO se reordenan ni se destacan por oferta) se agrupan en UNA
    sola línea de texto por subgrupo en vez de una píldora por ítem — con
    CVs de muchas skills (ej. Marcos: 12 core + 2 idiomas + 5 competencias)
    19 píldoras individuales, cada una con su propia mini-tabla, pesan
    demasiado en vertical y desbordan a una segunda página. Encontrado y
    arreglado el 11/08/2026."""
    core: list[dict] = []
    groups: list[dict] = []
    current_label: str | None = None
    current_items: list[str] = []

    def flush():
        if current_label is not None:
            groups.append({"label": current_label, "text": ", ".join(current_items)})

    pill_rank = 0
    for skill in skills:
        norm = _strip_accents(skill).upper().strip()
        if norm in _SKILL_SUBHEADER_KEYWORDS:
            flush()
            current_label = skill
            current_items = []
            continue
        if current_label is None:
            pill_rank += 1
            core.append({"text": skill, "type": "pill_gold" if pill_rank <= top_k else "pill_gray"})
        else:
            current_items.append(skill)
    flush()

    return {"core": core, "groups": groups}


def _set_compact_styles(doc):
    # Nota: este espaciado se aflojó (11/08/2026) — se había dejado muy
    # apretado pensando en el diseño viejo de una columna, que desbordaba
    # a 3 páginas. Con las dos columnas actuales sobraba mucho aire abajo
    # de la página con esos valores tan chicos.
    # line_spacing va SIEMPRE en 1.0 — probado en 1.05 y ese pequeño extra
    # multiplicado por las ~70 líneas de un CV cargado como el de Marcos
    # alcanza para tirarlo a una segunda página él solo.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.0

    bullets = doc.styles["List Bullet"]
    bullets.font.name = "Calibri"
    bullets.font.size = Pt(9.5)
    bullets.paragraph_format.space_before = Pt(0)
    bullets.paragraph_format.space_after = Pt(2)
    bullets.paragraph_format.line_spacing = 1.0


def _add_header_band(doc, name: str, tagline: str, contact_lines: list[str], content_width):
    table = doc.add_table(rows=1, cols=2)
    _remove_table_borders(table)
    name_w = Emu(int(content_width * 0.62))
    contact_w = Emu(int(content_width) - int(name_w))
    _set_col_widths(table, [name_w, contact_w])

    left, right = table.rows[0].cells
    for cell in (left, right):
        _shade_cell(cell, NAVY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    name_p = _next_para(left)
    _tight(name_p, space_after=2)
    name_run = name_p.add_run(name)
    name_run.bold = True
    name_run.italic = True
    name_run.font.size = Pt(25)
    name_run.font.name = "Cambria"
    name_run.font.color.rgb = _hex_to_rgb(WHITE)

    if tagline:
        tag_p = left.add_paragraph()
        _tight(tag_p, space_after=2)
        tag_run = tag_p.add_run(tagline)
        tag_run.font.size = Pt(9)
        tag_run.font.color.rgb = _hex_to_rgb(GRAY_ON_NAVY)

    for i, line in enumerate(contact_lines):
        p = _next_para(right) if i == 0 else right.add_paragraph()
        _tight(p, space_after=1)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(_contact_icon(line) + line)
        run.font.size = Pt(9)
        run.font.color.rgb = _hex_to_rgb(WHITE)


def _add_accent_bar(doc, content_width):
    table = doc.add_table(rows=1, cols=1)
    _remove_table_borders(table)
    _set_col_widths(table, [Emu(int(content_width))])
    row = table.rows[0]
    row.height = Pt(6)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell = row.cells[0]
    _shade_cell(cell, GOLD)
    p = cell.paragraphs[0]
    _tight(p, space_before=0, space_after=0)
    run = p.add_run(" ")
    run.font.size = Pt(2)


def _add_section_heading(cell, text: str):
    p = _next_para(cell)
    _tight(p, space_before=10, space_after=3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = _pt(10.5)
    run.font.color.rgb = _hex_to_rgb(NAVY)
    # Ojo: NO usar letter-spacing acá. Probado con 1pt extra por letra y en
    # columnas angostas Word corta la PALABRA a la mitad al hacer wrap (ej.
    # "PROYE CTOS DESTACADOS", "SOB RE M Í") — se ve roto, peor que sin nada.
    _set_paragraph_bottom_border(p, NAVY)


def _render_section(cell, sec: dict):
    _add_section_heading(cell, sec["heading"])
    kind = sec["kind"]
    content = sec["content"]

    if kind in ("perfil", "generic"):
        for text in content:
            p = cell.add_paragraph()
            _tight(p, space_after=4)
            run = p.add_run(text)
            run.font.size = _pt(9.5)
            run.font.color.rgb = _hex_to_rgb(GRAY_TEXT)

    elif kind == "proyectos":
        for item in content:
            p = cell.add_paragraph()
            _tight(p, space_before=(6 if item["is_title"] else 0), space_after=2)
            run = p.add_run(item["text"])
            run.font.size = _pt(9.5)
            if item["is_title"]:
                run.bold = True
                run.font.color.rgb = _hex_to_rgb(NAVY)
            else:
                run.font.color.rgb = _hex_to_rgb(GRAY_TEXT)

    elif kind == "experiencia":
        for block in content:
            p = cell.add_paragraph()
            _tight(p, space_before=7, space_after=0)
            run = p.add_run(block["company"])
            run.bold = True
            run.font.size = _pt(9.5)
            run.font.color.rgb = _hex_to_rgb(NAVY)
            for meta in block["meta_lines"]:
                mp = cell.add_paragraph()
                _tight(mp, space_after=1)
                mr = mp.add_run(meta)
                mr.italic = True
                mr.font.size = _pt(9)
                mr.font.color.rgb = _hex_to_rgb("595959")
            for bullet in block["bullets"]:
                bp = cell.add_paragraph(style="List Bullet")
                _tight(bp, space_after=1)
                br = bp.add_run(bullet)
                br.font.size = _pt(9.5)
                br.font.color.rgb = _hex_to_rgb(GRAY_TEXT)

    elif kind == "habilidades":
        pill_items = [
            (
                item["text"],
                GOLD if item["type"] == "pill_gold" else GRAY_LIGHT,
                NAVY if item["type"] == "pill_gold" else GRAY_TEXT,
                item["type"] == "pill_gold",
            )
            for item in content["core"]
        ]
        _add_skill_pill_grid(cell, pill_items, cols=2)

        for group in content["groups"]:
            label_p = cell.add_paragraph()
            _tight(label_p, space_before=4, space_after=1)
            label_run = label_p.add_run(group["label"])
            label_run.bold = True
            label_run.font.size = _pt(9.5)
            label_run.font.color.rgb = _hex_to_rgb(NAVY)

            text_p = cell.add_paragraph()
            _tight(text_p, space_after=2)
            text_run = text_p.add_run(group["text"])
            text_run.font.size = _pt(9.5)
            text_run.font.color.rgb = _hex_to_rgb(GRAY_TEXT)


def render_cv_docx(cv_data: dict, output_path: Path):
    """Genera el .docx con el template visual (único disponible por ahora,
    reconstruido a partir de la descripción del diseño real de Marcos en
    Canva). Genérico: sirve para cualquier CV, no sólo el de Marcos —
    sólo depende de la forma de cv_data, no de nada de Groq/adaptación.

    cv_data = {
        "name": str,
        "tagline": str,
        "contact_lines": [str, ...],
        "left_sections":  [{"heading": str, "kind": "perfil"|"proyectos"|"experiencia"|"generic", "content": ...}],
        "right_sections": [{"heading": str, "kind": "habilidades"|"generic", "content": ...}],
    }
    """
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    content_width = section.page_width - section.left_margin - section.right_margin

    _set_scale(1.0, 1.0)  # el header/barra dorada siempre van a tamaño fijo, no escalado
    _set_compact_styles(doc)
    _add_header_band(doc, cv_data["name"], cv_data.get("tagline", ""), cv_data.get("contact_lines", []), content_width)
    _add_accent_bar(doc, content_width)

    # El bloque "Oferta / Aplicar a esta oferta" YA NO va dentro del CV
    # (a pedido, 11/08/2026 — Marcos quiere el CV 100% fiel al diseño
    # original). Ese dato ahora vive en una hoja de portada aparte, ver
    # render_offers_cover().
    _set_scale(*_pick_scales(cv_data))

    body = doc.add_table(rows=1, cols=2)
    _remove_table_borders(body)
    left_w = Emu(int(content_width * 0.62))
    right_w = Emu(int(content_width) - int(left_w))
    _set_col_widths(body, [left_w, right_w])
    left_cell, right_cell = body.rows[0].cells
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    for sec in cv_data.get("left_sections", []):
        _render_section(left_cell, sec)
    for sec in cv_data.get("right_sections", []):
        _render_section(right_cell, sec)

    # Nota: a diferencia de una celda, agregar una tabla al DOCUMENTO no
    # inserta un párrafo automático después (comprobado — doc.paragraphs
    # queda vacío si el documento termina en tablas, y a Word no le
    # molesta). Antes esta función "tighteaba" doc.paragraphs[-1] pensando
    # que era un párrafo auto-insertado tras la tabla del cuerpo; en
    # realidad era el párrafo de "Aplicar a esta oferta" (ya sacado del
    # CV), así que esa línea no hacía lo que el comentario decía. Sacada.

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def _short_description(text: str, max_chars: int = 220) -> str:
    """Extracto LITERAL de la descripción real de la oferta — nada de
    reformular con IA acá, es justo el dato que tiene que ser 100% fiel a
    lo que dice el aviso. Sacamos tags HTML si los trae (Jobicy/Himalayas
    devuelven HTML), decodificamos entidades HTML (&amp; -> &, etc. —
    quedaban literales, encontrado 11/08/2026) y cortamos en el último
    espacio para no partir una palabra a la mitad."""
    if not text:
        return ""
    clean = re.sub(r"<[^>]+>", " ", text)
    clean = html.unescape(clean)
    clean = re.sub(r"\s+", " ", clean).strip()
    if len(clean) <= max_chars:
        return clean
    truncated = clean[:max_chars]
    last_space = truncated.rfind(" ")
    if last_space > 20:
        truncated = truncated[:last_space]
    return truncated + "…"


def render_offers_cover(jobs: list[dict], output_path: Path):
    """Genera UN solo documento con los datos de todas las ofertas: título,
    empresa, un extracto corto (real, no inventado) de la descripción, y
    el link para aplicar. Es la "hoja de portada" que acompaña a los CVs
    adaptados — separada del CV en sí (a pedido, 11/08/2026: el CV tiene
    que quedar 100% fiel al diseño original, sin ningún agregado adentro).
    El número de cada entrada corresponde al número del archivo del CV."""
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(2)
    title_run = title_p.add_run("Resumen de ofertas")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = _hex_to_rgb(NAVY)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(16)
    sub_run = sub_p.add_run(
        "Cada CV adaptado (01, 02, ...) corresponde a la oferta con el mismo número acá abajo."
    )
    sub_run.italic = True
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = _hex_to_rgb("595959")

    for i, job in enumerate(jobs, 1):
        heading_p = doc.add_paragraph()
        heading_p.paragraph_format.space_before = Pt(10)
        heading_p.paragraph_format.space_after = Pt(2)
        heading_run = heading_p.add_run(f"{i:02d}. {job.get('title') or 'sin título'}")
        heading_run.bold = True
        heading_run.font.size = Pt(12)
        heading_run.font.color.rgb = _hex_to_rgb(NAVY)

        if job.get("company"):
            company_p = doc.add_paragraph()
            company_p.paragraph_format.space_after = Pt(3)
            company_run = company_p.add_run(job["company"])
            company_run.italic = True
            company_run.font.size = Pt(10.5)
            company_run.font.color.rgb = _hex_to_rgb("595959")

        desc = _short_description(job.get("description", ""))
        if desc:
            desc_p = doc.add_paragraph()
            desc_p.paragraph_format.space_after = Pt(3)
            desc_run = desc_p.add_run(desc)
            desc_run.font.size = Pt(10)
            desc_run.font.color.rgb = _hex_to_rgb(GRAY_TEXT)

        if job.get("url"):
            link_p = doc.add_paragraph()
            link_p.paragraph_format.space_after = Pt(4)
            link_run = link_p.add_run(f"Aplicar: {job['url']}")
            link_run.bold = True
            link_run.font.size = Pt(10)
            link_run.font.color.rgb = _hex_to_rgb(NAVY)

        sep_p = doc.add_paragraph()
        sep_p.paragraph_format.space_after = Pt(2)
        _set_paragraph_bottom_border(sep_p, "CCCCCC", size=4)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def build_render_data(header_lines: list[str], sections: list[dict], adapted: dict) -> dict:
    """Adapta las estructuras internas del pipeline (header_lines/sections/
    adapted) al dict genérico que espera render_cv_docx — así el layout
    visual queda desacoplado de la lógica de adaptación con IA.

    El dato de la oferta (título/empresa/link) YA NO pasa por acá — desde
    el 11/08/2026 vive aparte, en la hoja de portada (render_offers_cover),
    no dentro del CV."""
    name = header_lines[0] if header_lines else "CV"
    tagline = header_lines[1] if len(header_lines) > 1 else ""
    contact_lines = header_lines[2:]

    left_sections: list[dict] = []
    right_sections: list[dict] = []

    for sec in sections:
        heading = sec["heading"]
        kind = _section_kind(heading)

        if kind == "perfil":
            text = adapted.get("perfil") or "\n".join(_rejoin_wrapped_lines(sec["body_lines"]))
            content = [p for p in text.split("\n") if p.strip()]
            left_sections.append({"heading": heading, "kind": "perfil", "content": content})

        elif kind == "proyectos":
            text = adapted.get("proyectos") or "\n".join(_rejoin_wrapped_lines(sec["body_lines"]))
            paras = [p for p in text.split("\n") if p.strip()]
            content = [{"text": p, "is_title": _looks_like_title(p)} for p in paras]
            left_sections.append({"heading": heading, "kind": "proyectos", "content": content})

        elif kind == "experiencia":
            blocks = sec.get("_blocks") or parse_experience_blocks(sec["body_lines"])
            exp_adapted = adapted.get("experiencia", {})
            content = [
                {
                    "company": b["company"],
                    "meta_lines": _rejoin_wrapped_lines(b["meta_lines"]),
                    "bullets": exp_adapted.get(b["company"], b["bullets"]),
                }
                for b in blocks
            ]
            left_sections.append({"heading": heading, "kind": "experiencia", "content": content})

        elif kind == "habilidades":
            skills = adapted.get("habilidades") or [l.strip() for l in sec["body_lines"] if l.strip()]
            right_sections.append({"heading": heading, "kind": "habilidades", "content": _build_skill_pills(skills)})

        else:
            paras = _rejoin_wrapped_lines(sec["body_lines"])
            # Educación acompaña a Habilidades en la columna derecha (angosta);
            # el resto (Sobre mí, etc.) va a la izquierda con el cuerpo principal.
            is_educacion = "EDUCAC" in _strip_accents(heading).upper() or "EDUCATION" in _strip_accents(heading).upper()
            target = right_sections if is_educacion else left_sections
            target.append({"heading": heading, "kind": "generic", "content": paras})

    return {
        "name": name,
        "tagline": tagline,
        "contact_lines": contact_lines,
        "left_sections": left_sections,
        "right_sections": right_sections,
    }


# ── Selección de ofertas a adaptar ───────────────────────────────
def load_top_jobs(cv_path: str, top_n: int) -> list[dict]:
    """Recalcula el prefiltro por embeddings en el momento (rápido, local,
    NO usa el modelo grande de Groq) para asegurar datos frescos — no
    reusa archivos de corridas previas, que pueden tener bugs ya
    arreglados (ej. company='Postular') o estar desactualizados."""
    from pipeline import collect_jobs
    from semantic_match import prefilter_by_embeddings

    cv_text = extract_cv_text(cv_path)
    print("\nClasificando CV con IA (modelo chico)...")
    profile = classify_cv(cv_text)
    print(f"  Área: {profile['area']} | is_tech: {profile['is_tech']}")

    jobs = collect_jobs(profile["search_queries"], profile.get("search_queries_en", []), profile["is_tech"])
    if not jobs:
        return []

    finalists = prefilter_by_embeddings(cv_text, jobs, top_k=top_n)
    return finalists


# ── Main ──────────────────────────────────────────────────────────
def main(cv_path: str, top_n: int = 10):
    print("=" * 60)
    print("  CV TAILOR — CVs adaptados por oferta (sin inventar nada)")
    print("=" * 60)

    cv_text = extract_cv_text(cv_path)
    header_lines, sections = parse_cv(cv_text)
    print(f"\nSecciones detectadas: {[s['heading'] for s in sections]}")

    jobs = load_top_jobs(cv_path, top_n)
    if not jobs:
        print("\n[!] No se encontraron ofertas para adaptar.")
        return

    slug = Path(cv_path).stem
    out_dir = Path("data") / "cvs_adaptados" / slug

    print(f"\nGenerando hasta {min(top_n, len(jobs))} CV(s) adaptado(s)...")
    generated = 0

    for i, job in enumerate(jobs[:top_n], 1):
        title = job.get("title", "sin título")
        company = job.get("company", "")
        print(f"\n[{i}/{min(top_n, len(jobs))}] {title} — {company}")

        if not job.get("description") and job.get("source") == "computrabajo" and job.get("url"):
            print("    Trayendo descripción completa del aviso...")
            job["description"] = fetch_full_description(job["url"])
            time.sleep(random.uniform(1.0, 2.0))  # no golpear el sitio sin pausa

        print("    Adaptando con IA (modelo chico, sin tocar el cupo agotado)...")
        adapted = adapt_for_job(cv_text, sections, job)

        safe_title = _safe_filename(title) or "oferta"
        safe_company = _safe_filename(company)
        filename = f"{i:02d}_{safe_title}"
        if safe_company:
            filename += f"_{safe_company}"
        output_path = out_dir / f"{filename}.docx"

        cv_data = build_render_data(header_lines, sections, adapted)
        render_cv_docx(cv_data, output_path)
        print(f"    -> {output_path}")
        generated += 1

    cover_path = out_dir / "00_Resumen_Ofertas.docx"
    render_offers_cover(jobs[:top_n], cover_path)
    print(f"\n✓ Hoja de portada (título/empresa/descripción/link, para las {min(top_n, len(jobs))} ofertas): {cover_path}")

    print(f"\n{'=' * 60}")
    print(f"✓ {generated} CV(s) adaptado(s) en: {out_dir.resolve()}")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("cv_path", nargs="?", default="test_cvs/cv_marcos.txt")
    parser.add_argument("--top", type=int, default=10)
    args = parser.parse_args()
    main(args.cv_path, args.top)
